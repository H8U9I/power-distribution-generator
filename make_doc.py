# -*- coding: utf-8 -*-
"""生成《配电图自动生成器 软件使用说明.docx》并附一张示意系统图。"""
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REL = r"H:\配电图生成器\release"
os.makedirs(REL, exist_ok=True)

# ---------- 1. 画一张示意系统图 (PNG) ----------
FONT_PATH = r"C:\Windows\Fonts\msyh.ttc"
try:
    fbig = ImageFont.truetype(FONT_PATH, 22)
    fmid = ImageFont.truetype(FONT_PATH, 16)
    fsm  = ImageFont.truetype(FONT_PATH, 13)
except Exception:
    fbig = fmid = fsm = ImageFont.load_default()

W, H = 1000, 760
img = Image.new("RGB", (W, H), "white")
d = ImageDraw.Draw(img)

def txt(x, y, s, f=fmid, fill="black"):
    d.text((x, y), s, font=f, fill=fill)

# 标题
txt(250, 14, "示例配电箱系统图（示意）", fbig, "black")
txt(330, 46, "三相 380V / 单相 220V  总开关 C63A  进线 BV-25", fsm, "#555555")

# 母线
bus_y = 90
d.line([(60, bus_y), (940, bus_y)], fill="black", width=4)
for i, (lab, x) in enumerate([("L1", 130), ("L2", 330), ("L3", 530), ("N", 730), ("PE", 870)]):
    txt(x, bus_y - 28, lab, fmid, "black")
# 三相不平衡标注
txt(700, bus_y + 8, "三相不平衡 ≤ 15%", fsm, "#006600")

# 总开关
d.rectangle([60, bus_y + 30, 160, bus_y + 90], outline="black", width=2)
txt(72, bus_y + 46, "总开关", fsm)
txt(72, bus_y + 66, "C63A", fsm)
d.line([(160, bus_y + 60), (220, bus_y + 60)], fill="black", width=2)

# 回路列表
loops = [
    ("照明回路", "C10A", "BV-2.5", False),
    ("普通插座", "C16A", "BV-2.5", True),
    ("挂机空调1", "C16A", "BV-2.5", False),
    ("挂机空调2", "C16A", "BV-2.5", False),
    ("柜机空调", "C20A", "BV-4", False),
    ("厨房插座", "C20A", "BV-4", True),
    ("卫生间插座", "C16A", "BV-2.5", True),
    ("热水器", "C16A", "BV-4", True),
]
x0 = 220
col_w = 92
y0 = bus_y + 60
for idx, (name, brk, wire, rcd) in enumerate(loops):
    cx = x0 + idx * col_w
    # 竖线从母线引下
    d.line([(cx, bus_y + 4), (cx, y0)], fill="black", width=2)
    # 断路器框
    by = y0
    d.rectangle([cx - 34, by, cx + 34, by + 56], outline="black", width=2)
    txt(cx - 30, by + 4, name, fsm)
    txt(cx - 30, by + 22, brk, fsm)
    txt(cx - 30, by + 38, wire, fsm)
    if rcd:
        txt(cx - 30, by + 54, "带漏保", fsm, "#CC0000")
    # 出线
    d.line([(cx, by + 56), (cx, by + 80)], fill="black", width=2)
    txt(cx - 30, by + 84, "→设备", fsm, "#555555")

txt(60, H - 50, "说明：示意仅展示布局样式，实际回路/线缆/空开由软件按输入自动计算生成。", fsm, "#888888")
png_path = os.path.join(REL, "_示意图.png")
img.save(png_path, "PNG")
print("示意图已生成:", png_path)

# ---------- 2. 构建 docx ----------
doc = Document()
# 默认中文字体
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
try:
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
except Exception:
    pass

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, bold=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def numbered(text):
    doc.add_paragraph(text, style="List Number")

# 封面标题
t = doc.add_heading("配电图自动生成器", level=0)
para("软件使用说明  ·  v2.0", bold=True, size=13, color=(0x1F, 0x4E, 0x79))
para("适用系统：Windows 10 / 11（64 位）", size=9, color=(0x66, 0x66, 0x66))
doc.add_picture(png_path, width=Inches(6.3))
last = doc.paragraphs[-1]
last.alignment = WD_ALIGN_PARAGRAPH.CENTER
para("图 1  软件生成的配电箱系统图示意", size=9, color=(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER)

# 1. 软件简介
heading("一、软件简介", 1)
para("本软件用于根据「设备 + 用电量」自动绘制符合国家低压配电常规做法的配电箱系统图。"
     "支持三相电与三相平衡计算，按设备功率自动匹配线缆规格与回路断路器容量，"
     "可导出 SVG（矢量图）、DXF（CAD 通用格式）、DWG（AutoCAD 原生格式）及负荷计算书。")
para("特点：", bold=True)
bullet("可视化管理界面，以「添加设备」为核心，无需写代码、无需手画图纸。")
bullet("支持三相平衡分配，自动显示三相电流与不平衡度。")
bullet("内置防火等级与线缆类型配置（普通 / 阻燃 / 耐火 / 低烟无卤；BV / BVR / YJV 等）。")
bullet("回路分三类管理：插座回路、照明回路、设备回路，均支持层级子菜单。")
bullet("一键导出 CAD 可用的 DWG（需本机安装 AutoCAD，详见第三节）。")

# 2. 软件包内容
heading("二、软件包内容（解压后）", 1)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
hdr[0].text = "文件 / 文件夹"; hdr[1].text = "说明"; hdr[2].text = "是否必需"
rows = [
    ("配电图生成器.exe", "主程序，双击即可运行", "必需"),
    ("vc_redist.x64.exe", "Microsoft Visual C++ 运行库安装包（兜底用）", "按需"),
    ("软件使用说明.docx", "本说明文档", "参考"),
    ("依赖说明.txt", "在别的电脑上使用的依赖与注意事项", "参考"),
]
for r in rows:
    c = tbl.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v
para("导出文件默认生成在 exe 同目录下的 output\\ 文件夹中。", size=9, color=(0x66,0x66,0x66))

# 3. 在别的电脑上使用（依赖重点）
heading("三、在别的电脑上使用（重要）", 1)
para("本软件已打包为单文件 exe，Python 运行环境与所需库均已内置，"
     "因此「别的电脑上不需要安装 Python」。但有以下几点依赖需注意：", bold=True)

para("1. 是否需要安装 Python？——不需要。", bold=True)
bullet("exe 已将 Python 解释器、ezdxf、openpyxl、win32com 等全部打包进去，双击即用。")

para("2. 是否需要安装运行库？——多数情况不需要，极少数需补一个。", bold=True)
bullet("部分未安装过 Visual C++ 运行库的 Windows，打开 exe 可能报错「缺少 VCRUNTIME140.dll / MSVCP140.dll」。")
bullet("解决办法：双击运行包里的 vc_redist.x64.exe，一路「下一步」安装完成，再开 exe 即可。")
bullet("Windows 10 / 11 绝大多数已自带该运行库，通常直接能跑。")

para("3. 导出 DWG 是否需要 AutoCAD？——需要本机装有 AutoCAD。", bold=True)
bullet("DWG 是 AutoCAD 原生私有格式。软件通过调用本机已安装的 AutoCAD（COM 自动化）"
       "将内部生成的 DXF 另存为 DWG；开发机已用 AutoCAD 2014 实测通过，其他 AutoCAD 版本（2010 及以上）通常兼容。")
bullet("若目标电脑没有 AutoCAD：SVG / DXF / 计算书 导出完全正常，不受影响。")
bullet("DXF 是开放格式，可用中望 CAD、AutoCAD、LibreCAD 或在线 DXF 查看器打开，"
       "「文件→另存为」即可转成 dwg。")
bullet("无 AutoCAD 也可使用免费的「ODA File Converter」（需自行下载）："
       "把软件导出的 .dxf 拖入该程序，选择输出格式 DWG、版本 ACAD2013，即可转换。")

para("结论：把整个压缩包解压到任意 Windows 电脑，双击 exe 基本就能用；"
     "万一报缺运行库就装一下 vc_redist；要出 DWG 就确保有 AutoCAD（或走 DXF 中转）。",
     bold=True, color=(0x1F, 0x4E, 0x79))

# 4. 快速开始
heading("四、快速开始（5 步）", 1)
numbered("解压压缩包到任意目录（路径不要含中文空格过长的特殊字符也可，中文目录已验证可用）。")
numbered("双击「配电图生成器.exe」。程序会自动启动内置服务并打开默认浏览器。")
numbered("若浏览器未自动打开，手动访问地址栏中的 http://127.0.0.1:18800 。")
numbered("在左侧「项目设置」选择电压系统（三相/单相）、防火等级、线缆类型；点「载入示例」可先体验。")
numbered("按需在左侧添加设备与回路，右侧实时显示系统图；点顶部「导出 DWG / DXF / SVG / 计算书」下载成果。")
para("提示：软件关闭浏览器或退出时，后台服务会随 exe 退出；若需再次使用，重新双击 exe 即可。",
     size=9, color=(0x66,0x66,0x66))

# 5. 操作详解
heading("五、界面与操作详解", 1)
heading("5.1 项目设置", 2)
bullet("电压系统：三相 380V（带三相平衡）或 单相 220V。")
bullet("防火等级：普通 / 阻燃C / 阻燃B / 耐火 / 低烟无卤，影响线缆载流量降容系数（耐火×0.85 等）。")
bullet("线缆类型：BV / BVR / YJV / NH-BV 等，决定导线前缀与用途。")

heading("5.2 添加设备与回路（以添加设备为主）", 2)
para("左侧为层级树：配电箱 → 回路类型 → 分路 → 设备。三类回路：")
bullet("插座回路：可填写插座总数与分配路数，常用插座与单独电器分开设路，均带漏电保护（RCD）。")
bullet("照明回路：可配置是否分多路及分路数量，照明回路一般不分设 RCD。")
bullet("设备回路：独立设备或设备组，内含设备列表子菜单，可逐一添加该回路下的设备（支持三相设备）。")
para("每个设备需填写：名称、功率（W）、数量。系统按功率自动匹配线缆规格与回路空开容量。",
     bold=True)
para("每种回路类型均提供层级子菜单，支持分路数量配置与设备明细管理；高级项可覆盖需用系数 kx、"
     "功率因数 cosφ、是否带 RCD、是否三相设备。", size=9, color=(0x66,0x66,0x66))

heading("5.3 实时预览", 2)
bullet("右侧画布实时显示配电箱系统图：母线、总开关、各回路断路器、线缆与空开标注、RCD 标记。")
bullet("三相模式下实时显示三相电流与不平衡度，便于检查负载分配是否均衡。")

heading("5.4 导出", 2)
bullet("导出 SVG：矢量图，浏览器即可打开，放大不模糊，适合汇报与插入文档。")
bullet("导出 DXF：CAD 通用交换格式，中望 / AutoCAD / LibreCAD 直接打开编辑。")
bullet("导出 DWG：调用本机 AutoCAD 自动另存为原生 dwg（需装有 AutoCAD）。")
bullet("导出计算书：Markdown 文件，含每回路功率、计算电流、线缆、空开明细。")

# 6. 计算规则
heading("六、自动匹配的计算规则（简化国标）", 1)
bullet("电压：单相 220V / 三相 380V；功率因数 cosφ 默认 0.9。")
bullet("需用系数 kx：插座 0.6、照明 0.9、专用设备 1.0（可在高级项覆盖）。")
bullet("线缆选型：按铜芯载流量表选，防火等级做降容（普通×1.0 / 阻燃×0.9 / 耐火×0.85 / 低烟无卤×0.85）。")
bullet("空开（C 型微断）：按计算电流向上取整；插座不低于 C16A、照明不低于 C10A、"
       "空调等电机类上浮至 C16A 起步。")
bullet("最小线径：住宅固定布线不低于 2.5mm²；大功率 / 柜机 / 热水器用 4mm²。")
bullet("三相平衡：单相回路按计算电流降序贪心填入当前电流最小的相，降低不平衡度。")

# 7. 常见问题
heading("七、常见问题（FAQ）", 1)
faqs = [
    ("双击 exe 没反应 / 报缺 dll？", "运行包内 vc_redist.x64.exe 安装运行库后重试；确认系统是 64 位 Windows。"),
    ("浏览器没自动打开？", "手动打开浏览器访问 http://127.0.0.1:18800 ；若提示无法访问，确认 exe 正在运行。"),
    ("点「导出 DWG」失败？", "确认本机装有 AutoCAD；无 AutoCAD 时请用「导出 DXF」再于 CAD 中另存为 dwg，或用 ODA File Converter 中转。"),
    ("能导入 Excel 设备清单吗？", "当前版本在界面手动添加；后端已支持读取 xlsx，如需批量导入可在后续版本加入（可联系开发者）。"),
    ("出图能直接拿去施工吗？", "不能。本软件为辅助设计工具，正式施工图须由注册电气工程师按 GB 50054 / GB 51348 校核。"),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run("Q：" + q); r.bold = True
    doc.add_paragraph("A：" + a)

# 8. 免责声明
heading("八、免责声明", 1)
para("本软件生成的图纸与计算书仅供方案构思、教学演示与内部沟通参考，"
     "不得直接作为正式施工依据。任何实际工程应用必须由具备资质的专业电气工程师"
     "依据国家现行规范（GB 50054《低压配电设计规范》、GB 51348《民用建筑电气设计标准》等）"
     "进行设计、校核与签章。因使用本软件产生的任何后果，开发者不承担责任。",
     size=9, color=(0x66,0x66,0x66))

out = os.path.join(REL, "软件使用说明.docx")
doc.save(out)
print("使用说明已生成:", out)
print("大小:", os.path.getsize(out), "字节")
